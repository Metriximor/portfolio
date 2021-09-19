import spotipy
import operator
from spotipy.oauth2 import SpotifyClientCredentials
from lyricsgenius import Genius
from docx import Document

def get_song_list(spotipy_playlist): 
    songs_list = []
    songs = spotipy_playlist['items']
    for song in songs:
        song_data = {}
        song_data['artist'] = song['track']['artists'][0]['name']
        song_data['name'] = song['track']['name']
        songs_list.append(song_data)
    return songs_list

auth_manager = SpotifyClientCredentials(client_id=CLIENT_ID,client_secret=CLIENT_SECRET)
sp = spotipy.Spotify(auth_manager=auth_manager)
genius = Genius(GENIUS_ACCESS_TOKEN)

playlist = sp.playlist_items(PLAYLIST_URL)
songs = get_song_list(playlist)
songs = sorted(songs, key= operator.itemgetter('artist', 'name'))

# Fetching Lyrics
for song in songs:
    try:
        lyrics = genius.search_song(song['name'], artist=song['artist'])
        song['lyrics'] = lyrics.lyrics
    except:
        song['lyrics'] = "Não encontrado"

# Writing in Word Doc
doc = Document()
for song in songs:
    p = doc.add_paragraph()
    r = p.add_run(f"{song['artist']} - {song['name']}")
    r.bold = True
    doc.add_paragraph(f"{song['lyrics']}")
    doc.add_page_break()
doc.save("lyrics.docx")
        
